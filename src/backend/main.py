import os
import io
import json
from pathlib import Path
from typing import Any, Dict, List
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_classic.chains import create_retrieval_chain
from langchain_classic.chains.combine_documents import (
    create_stuff_documents_chain,
)
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from pathlib import Path
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

import hashlib

HASHES_FILE = "backend/processed_hashes.json"

load_dotenv()

app = FastAPI(title="Simple React Backend")


def compute_content_hash(content):
    return hashlib.sha256(content).hexdigest()



def load_processed_hashes():
    if os.path.exists(HASHES_FILE):
        with open(HASHES_FILE, "r") as f:
            return set(json.load(f))
    return set()

def save_processed_hashes(hashes):
    os.makedirs(os.path.dirname(HASHES_FILE), exist_ok=True)
    with open(HASHES_FILE, "w") as f:
        json.dump(list(hashes), f)

@app.post("/api/research")
async def do_rag(data: Dict[str, str]):
    faiss_path = "backend/faiss_index"

    embeddings = OpenAIEmbeddings()

    db = FAISS.load_local(folder_path=faiss_path, embeddings=embeddings,
                          allow_dangerous_deserialization=True)

    retriever = db.as_retriever(search_kwargs={"k": 3})

    template = """You are the best documentation reader and explanator.
    Use the context to answer the research query from the user. 
    Use markdown in your answer. For example bullet points or titles.
    If you do not have enough information, answer that you need more documentation. 
    
    Context: {context}
    
    Research query: {input}
    """
    prompt = ChatPromptTemplate.from_template(template=template)
    llm = ChatOpenAI(model="gpt-3.5-turbo-0125", temperature=0)
    combine_docs_chain = create_stuff_documents_chain(llm, prompt)

    qa_chain = create_retrieval_chain(retriever, combine_docs_chain)


    async def generate(query):
        async for chunk in qa_chain.astream({"input": query['request']}):
            text = chunk.get("answer", "")
            yield text

    to_return = StreamingResponse(
        content=generate(data),
        media_type="text/markdown",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive"
        }
    )
    return to_return


@app.post("/api/sources")
async def upload_files(files: List[UploadFile] = File(...)):
    processed_hashes = load_processed_hashes()
    to_return = []
    docs = []
    for f in files:
        safe_name = f.filename.replace("\\", "/").split("/")[-1]
        if not safe_name:
            raise HTTPException(status_code=400, detail="Empty filename")
        extension = f.filename.lower().split('.')[-1]
        content_bytes = await f.read()
        file_hash = compute_content_hash(content_bytes)

        if file_hash not in processed_hashes:
            processed_hashes.add(file_hash)
            ff = ""
            if extension not in ["pdf", "doc", "docx"]:
                ff = content_bytes.decode("utf-8",  errors="ignore")
            elif extension == "pdf":
                reader = PdfReader(io.BytesIO(content_bytes))
                ff = ""
                for page in reader.pages:
                    ff += page.extract_text() or ""
            elif extension in ['doc', 'docx']:
                doc = DocxDocument(io.BytesIO(content_bytes))
                ff = "\n".join([para.text for para in doc.paragraphs])
            doc = Document(
                page_content=ff,
                metadata={"source": safe_name}
            )
            docs.append(doc)

        to_return.append({
            "name": safe_name,
            "size": f.size,
            "type": f.content_type
        })
    save_processed_hashes(processed_hashes)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
    )
    if len(docs) > 0:
        split_docs = text_splitter.split_documents(docs)
        embeddings = OpenAIEmbeddings()

        faiss_index_path = "backend/faiss_index"
        if os.path.exists(faiss_index_path):
            db = FAISS.load_local(
                folder_path=faiss_index_path,
                embeddings=embeddings,
                allow_dangerous_deserialization=True
            )
            db.add_documents(split_docs)
        else:
            db = FAISS.from_documents(split_docs, embeddings)

        db.save_local(faiss_index_path)

    return {"uploaded": to_return}