import json
import os
import io
import hashlib
from docx import Document as DocxDocument
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from constants import HASHES_FILE, CHUNK_SIZE, CHUNK_OVERLAP, FAISS_PATH


def compute_content_hash(content):
    return hashlib.sha256(content).hexdigest()


def load_processed_hashes():
    if os.path.exists(HASHES_FILE):
        with open(HASHES_FILE, "r") as f:
            return set(json.load(f))
    return set()


def save_processed_hashes(hashes):
    os.makedirs(os.path.dirname(HASHES_FILE), exist_ok=True)
    with open(HASHES_FILE, "w") as f:
        json.dump(list(hashes), f)


async def add_documents_to_index(docs):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    split_docs = text_splitter.split_documents(docs)
    embeddings = OpenAIEmbeddings()

    if os.path.exists(FAISS_PATH):
        db = FAISS.load_local(
            folder_path=FAISS_PATH,
            embeddings=embeddings,
            allow_dangerous_deserialization=True
        )
        db.add_documents(split_docs)
    else:
        db = FAISS.from_documents(split_docs, embeddings)

    db.save_local(FAISS_PATH)


async def add_file_to_the_documentation(content_bytes,
                                        docs,
                                        extension,
                                        file_hash,
                                        processed_hashes,
                                        safe_name):
    processed_hashes.add(file_hash)
    ff = ""
    if extension not in ["pdf", "doc", "docx"]:
        ff = content_bytes.decode("utf-8", errors="ignore")
    elif extension == "pdf":
        reader = PdfReader(io.BytesIO(content_bytes))
        ff = ""
        for page in reader.pages:
            ff += page.extract_text() or ""
    elif extension in ['doc', 'docx']:
        doc = DocxDocument(io.BytesIO(content_bytes))
        ff = "\n".join([para.text for para in doc.paragraphs])
    doc = Document(
        page_content=ff,
        metadata={"source": safe_name}
    )
    docs.append(doc)
